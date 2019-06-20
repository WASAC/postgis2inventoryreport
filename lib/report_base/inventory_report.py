import os
import shutil
import datetime
from docx import Document
from lib.report_base.report_base import ReportBase
from lib.report_base.assets_base.wss import WssList
from lib.report_base.assets_base.summary import Summary
from lib.report_base.assets_base.chamber import Chambers
from lib.report_base.assets_base.reservoir import Reservoirs
from lib.report_base.assets_base.pumping_station import PumpingStations
from lib.report_base.assets_base.watersource import WaterSources
from lib.report_base.assets_base.waterconnection import WaterConnections


class InventoryReport(ReportBase):
    def __init__(self, db, district, tmp_file_path, main_directory):
        self.db = db
        self.district = district
        self.tmp_file_path = tmp_file_path
        self.main_directory = main_directory
        current_date=datetime.datetime.now()
        self.month_year = current_date.strftime('%B %Y')

    def create(self):
        if not os.path.exists(self.main_directory):
            os.makedirs(self.main_directory, exist_ok=True)

        new_file_path = "/".join(
            [self.main_directory, "{0}_Inventory Data for {1} District.docx".format(self.district.dist_id, self.district.district)])
        shutil.copy2(self.tmp_file_path, new_file_path)
        dist_image_dir = "/".join([self.main_directory, "images", str(self.district.dist_id)])

        doc = Document(new_file_path)

        values = {"%district%": self.district.district.upper(), "%monthyear%": self.month_year}
        self.docx_replace(doc, values)

        doc.add_heading('Rural Water Suppy Systems in {0} District'.format(self.district.district), level=1)
        doc.add_heading('About Summary of District', level=2)
        doc.add_paragraph('Please describe the WSS here. ')
        doc.add_heading('Map of the entire systems', level=2)
        dist_image = "/".join([dist_image_dir, "{0}.png".format(str(self.district.dist_id))])
        if os.path.exists(dist_image):
            self.add_image(doc, dist_image)
        else:
            self.add_temp_image(doc)
        doc.add_page_break()

        wss_list_obj = WssList(self.district.dist_id)
        wss_list = wss_list_obj.create(self.db, doc)

        for wss_data in wss_list:
            doc.add_heading('{0} {1} WSS'.format(wss_data.wss_id, wss_data.wss_name), level=1)
            doc.add_heading('About Summary of {0} WSS'.format(wss_data.wss_name), level=2)
            if not wss_data.description:
                doc.add_paragraph('Description for the WSS is not yet registered in GIS database. ')
            else:
                doc.add_paragraph(wss_data.description)
            doc.add_heading('Map of the WSS', level=2)
            wss_image = "/".join([dist_image_dir, "{0}.png".format(str(wss_data.wss_id))])
            if os.path.exists(wss_image):
                self.add_image(doc, wss_image)
            else:
                self.add_temp_image(doc)
            doc.add_page_break()
            doc.add_heading('Assets Data', level=2)

            for assets_obj in [Summary(wss_data.wss_id), Chambers(wss_data.wss_id),
                               Reservoirs(wss_data.wss_id), PumpingStations(wss_data.wss_id),
                               WaterSources(wss_data.wss_id), WaterConnections(wss_data.wss_id)]:
                assets_obj.create(self.db, doc)
            doc.add_page_break()
        doc.save(new_file_path)
        del doc
        print("It created Inventory Report of {0} at folder: {1}".format(self.district.district, new_file_path))
